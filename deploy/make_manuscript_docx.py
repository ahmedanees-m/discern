"""Render the BMC submission manuscript as a submission-ready .docx.

BMC accepts DOC/DOCX/RTF/TeX for the main text and asks for double-line spacing with line and page
numbering, none of which markdown carries. This builds the Word file directly from the markdown so
the manuscript still has a single source.

Run:  python -m deploy.make_manuscript_docx SRC.md OUT.docx
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def convert(SRC, DST):
    doc = Document()

    # Double-line spacing and a plain serif face, set on Normal so everything inherits it.
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)

    sec = doc.sections[0]

    # Continuous line numbering, which BMC asks for and Word only exposes as a section property.
    sect_pr = sec._sectPr
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "continuous")
    ln.set(qn("w:distance"), "360")
    sect_pr.append(ln)

    # Page number, centred in the footer, as a PAGE field rather than literal text.
    foot = sec.footer.paragraphs[0]
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = foot.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for e in (begin, instr, end):
        run._r.append(e)

    INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`)")


    def add_runs(par, text):
        """Emit a paragraph's text, honouring bold, italic and inline code."""
        for part in INLINE.split(text):
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                par.add_run(part[2:-2]).bold = True
            elif part.startswith("*") and part.endswith("*"):
                par.add_run(part[1:-1]).italic = True
            elif part.startswith("`") and part.endswith("`"):
                r = par.add_run(part[1:-1])
                r.font.name = "Consolas"
                r.font.size = Pt(10)
            else:
                par.add_run(part)


    def is_table_row(s):
        return s.startswith("|") and s.endswith("|")


    lines = open(SRC, encoding="utf-8").read().split("\n")
    i, buf, n_tables = 0, [], 0


    def flush():
        if buf:
            add_runs(doc.add_paragraph(), " ".join(buf))
            buf.clear()


    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if is_table_row(stripped):
            flush()
            block = []
            while i < len(lines) and is_table_row(lines[i].strip()):
                block.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            block = [r for r in block if not all(set(c) <= set("-: ") for c in r)]
            if block:
                t = doc.add_table(rows=len(block), cols=len(block[0]))
                t.style = "Table Grid"
                for r, row in enumerate(block):
                    for c, cell in enumerate(row[:len(block[0])]):
                        cell_par = t.cell(r, c).paragraphs[0]
                        cell_par.paragraph_format.line_spacing = 1.0
                        add_runs(cell_par, cell)
                        if r == 0:
                            for rr in cell_par.runs:
                                rr.bold = True
                n_tables += 1
            continue

        if stripped.startswith("#"):
            flush()
            level = len(stripped) - len(stripped.lstrip("#"))
            doc.add_heading(stripped.lstrip("# ").strip(), level=min(level, 4))
        elif stripped == "---":
            flush()
        elif stripped.startswith("- "):
            flush()
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, stripped[2:])
        elif not stripped:
            flush()
        else:
            buf.append(stripped)
        i += 1
    flush()

    doc.save(DST)
    return DST



if __name__ == "__main__":
    print("wrote", convert(sys.argv[1], sys.argv[2]))
